import streamlit as st
from docx import Document
from docx.shared import Pt


def replace_text(doc, old_text, new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    text = inline[i].text.replace(old_text, new_text)
                    inline[i].text = text


def add_job_section(doc, job_info, language_section):
    paragraph = doc.paragraphs[language_section].insert_paragraph_before()

    def add_bold_label(text):
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'

    def add_normal_text(text):
        run = paragraph.add_run(text)
        run.bold = False
        run.font.size = Pt(11)
        run.font.name = 'Arial'

    add_bold_label("Name of Employer: ")
    add_normal_text(f"{job_info['name_of_employer']}\n")

    add_bold_label("Dates of Employment: ")
    add_normal_text(f"{job_info['dates_of_employment']}\n")

    add_bold_label("Job Title: ")
    add_normal_text(f"{job_info['job_title']}\n")

    add_bold_label("Project/Role Description: ")
    add_normal_text(f"{job_info['project_role_description']}\n")


def replace_text_in_docx(template_path, replacements, jobs):
    doc = Document(template_path)
    for old_text, new_text in replacements.items():
        replace_text(doc, old_text, new_text)

    language_section = None
    for i, paragraph in enumerate(doc.paragraphs):
        if 'Languages' in paragraph.text:
            language_section = i
            break

    if language_section is not None:
        for job_info in reversed(jobs):
            add_job_section(doc, job_info, language_section)

    return doc


st.title('Resume Updater')

name_surname = st.text_input('Name Surname', 'Murad Sofizade')
title = st.text_input('Title', 'CEO')
summary = st.text_area('Summary', 'I love to work with people and help them to achieve their goals')
skills = st.text_area('Skills', 'Teaching')
english_level = st.text_input('English level', 'Upper Intermediate')
education = st.text_area('Education', 'Master degree in computer science')
certifications = st.text_area('Certifications', 'PMP, Scrum Master, TOEFL, IELTS')


jobs = []
if 'jobs' not in st.session_state:
    st.session_state.jobs = []

add_job = st.button('Add Job')

if add_job:
    st.session_state.jobs.append({
        "name_of_employer": "",
        "dates_of_employment": "",
        "job_title": "",
        "project_role_description": ""
    })

for i, job in enumerate(st.session_state.jobs):
    st.session_state.jobs[i]["name_of_employer"] = st.text_input(f'Name of employer {i + 1}',
                                                                 key=f'name_of_employer{i + 1}')
    st.session_state.jobs[i]["dates_of_employment"] = st.text_input(f'Dates of employment {i + 1}',
                                                                    key=f'dates_of_employment{i + 1}')
    st.session_state.jobs[i]["job_title"] = st.text_input(f'Job title {i + 1}', key=f'job_title{i + 1}')
    st.session_state.jobs[i]["project_role_description"] = st.text_area(f'Project/Role description {i + 1}',
                                                                         key=f'project_role_description{i + 1}')


template_path = 'template.docx'
output_path = 'updated_resume.docx'


if st.button('Update Resume'):
    replacements = {
        "text1": name_surname,
        "text2": title,
        "text3": summary,
        "text4": skills,
        "eng1": english_level,
        "edu1": education,
        "cert1": certifications
    }

    doc = replace_text_in_docx(template_path, replacements, st.session_state.jobs)
    doc.save(output_path)
    st.success('Resume updated successfully!')

    with open(output_path, "rb") as file:
        st.download_button(label="Download Updated Resume", data=file, file_name="updated_resume.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


